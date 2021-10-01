# from selenium import webdriver
# from selenium.webdriver.common.by import By
# import time
# driver = webdriver.Chrome(executable_path='/Users/HP/Desktop/folders/others/chromedriver/chromedriver.exe')


# url=f"https://gre.magoosh.com/flashcards/vocabulary/decks"
# #url2=f'https://gre.magoosh.com/flashcards/vocabulary/common-words-{}'
# driver.get(url)
# xpath='/html/body/div/div/div[1]/div/a[2]'
# driver.find_element_by_xpath(xpath).click()
#        # /html/body/div/div/div/div[2]/div/div[1]/a[2]
# word='/html/body/div/div/div/div[1]/div/div[1]/a[1]/div/h1'
# #text1=driver.find_element_by_class_name('flashcard-word').text
# #print(f"i am text broo{text1}")



# #print(text1)
# mainlist=[]
# for i in range(0,51):
#     xpath2='/html/body/div/div/div/div[2]/div/div[1]/a[2]'
            
    
#     driver.find_element_by_class_name('flashcard').click()
#     temp_list=[]
#     text1=driver.find_element_by_class_name('flashcard-word').text
#     text2=driver.find_element_by_class_name('flashcard-text').text
#     text3=driver.find_element_by_class_name('flashcard-example').text
#     temp_list.append(text1)
#     temp_list.append(text2)
#     #temp_list.append(text3)
#     mainlist.append(temp_list)
#     idk='/html/body/div/div/div/div[1]/div/div[2]/a[2]'
         
    
#     driver.find_element_by_css_selector('body > div > div > div > div.flashcard-container.u-margin-V-m > div > div.back.card.flashcard-card > a.card-footer.card-footer-danger.text-center').click()
    
#     print(f"{i} {temp_list}")
#     time.sleep(2)
    
# j=0
# for i in mainlist:
#     print(f"{j+1}{i}")
#     j+=1
          
# # print(text3)


# # print(text2)
# # text4=driver.find_element_by_tag_name('strong').text
# #print(xvh)


















from selenium import webdriver
from selenium.webdriver.common.by import By
import time
driver = webdriver.Chrome(executable_path='/Users/HP/Desktop/folders/others/chromedriver/chromedriver.exe')

for k in range(2,22):
    try:
        url=f"https://gre.magoosh.com/flashcards/vocabulary/decks"
        #url2=f'https://gre.magoosh.com/flashcards/vocabulary/common-words-{}'
        driver.get(url)

        xpath=f"/html/body/div/div/div[{k}]/div/a[2]"

        driver.find_element_by_xpath(xpath).click()
               # /html/body/div/div/div/div[2]/div/div[1]/a[2]
        word='/html/body/div/div/div/div[1]/div/div[1]/a[1]/div/h1'
        #text1=driver.find_element_by_class_name('flashcard-word').text
        #print(f"i am text broo{text1}")

        z=1

        #print(text1)
        list1=[]
        while len(list1)<=51:
            xpath2='/html/body/div/div/div/div[2]/div/div[1]/a[2]'


            driver.find_element_by_class_name('flashcard').click()
            temp_list=[]
            text1=driver.find_element_by_class_name('flashcard-word').text
            text2=driver.find_element_by_class_name('flashcard-text').text
            text3=driver.find_element_by_class_name('flashcard-example').text
            temp_list.append(text1)
            temp_list.append(text2)
            #temp_list.append(text3)
            if temp_list not in list1:
                list1.append(temp_list)

                print(f"{z} {temp_list}")
                z+=1
            idk='/html/body/div/div/div/div[1]/div/div[2]/a[2]'


            driver.find_element_by_css_selector('body > div > div > div > div.flashcard-container.u-margin-V-m > div > div.back.card.flashcard-card > a.card-footer.card-footer-danger.text-center').click()

            #print(f"{i} {temp_list}")
            time.sleep(1)
        #m=[['hey','hey'],['hey','hey'],['fr','khgesrg']]
    except:
        pass
    from docx import Document
    #list1=[['hey','hey2'],['hey2','jdskyfgvbut'],['hey','hey2'],['hey2','jdskyfgvbut']]
    l=len(list1)
    list2=[]
    for i in range(0,l,2):
        li=[]
        li.append(list1[i][0])
        li.append(list1[i][1])
        li.append(list1[i+1][0])

        li.append(list1[i+1][1])

        list2.append(li)
    l=len(list2)
    doc = Document('magoosh.docx')
    doc.tables
    n=27
    l=l+27
    for i in range(n,l):
        doc.tables[0].cell(i, 0).text = list2[i][0]
        doc.tables[0].cell(i, 1).text = list2[i][1]
        doc.tables[0].cell(i, 2).text = list2[i][2]
        doc.tables[0].cell(i, 3).text = list2[i][3]
        doc.tables[0].add_row() #ADD ROW HERE
    doc.save("magoosh.docx")
    print(f"added {k}th flashcard")

